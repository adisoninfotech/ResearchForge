"""DOCX renderer from canonical manuscript (via python-docx)."""

from __future__ import annotations

import io

from app.services.export.canonical import CanonicalManuscript
from app.services.export.templates import get_template


def render_docx(manuscript: CanonicalManuscript) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    tpl = get_template(manuscript.template_id)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    fm = manuscript.front_matter
    title = doc.add_heading(fm.title or "Untitled", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if fm.authors:
        p = doc.add_paragraph(", ".join(a.name for a in fm.authors))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for aff in fm.affiliations:
        p = doc.add_paragraph(aff.name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True if p.runs else None

    if fm.abstract:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(fm.abstract)
    if fm.keywords:
        doc.add_paragraph("Keywords: " + ", ".join(fm.keywords))

    for section in manuscript.sections:
        if section.section_type in {"abstract", "keywords"}:
            continue
        doc.add_heading(section.title, level=1)
        for block in section.blocks:
            if block.type == "heading":
                doc.add_heading(block.text, level=min(block.level or 2, 3))
            elif block.type == "list":
                style_name = "List Number" if block.ordered else "List Bullet"
                for item in block.items:
                    doc.add_paragraph(item, style=style_name)
            elif block.type == "equation":
                doc.add_paragraph(f"Equation: {block.latex or block.text}")
            elif block.type == "figure":
                fig = next((f for f in manuscript.figures if f.id == block.figure_id), None)
                if fig:
                    doc.add_paragraph(f"[Figure {fig.number}]")
                    doc.add_paragraph(f"Figure {fig.number}: {fig.caption or fig.title}")
                else:
                    doc.add_paragraph(block.text)
            elif block.type == "table":
                tab = next((t for t in manuscript.tables if t.id == block.table_id), None)
                if tab and (tab.headers or tab.rows):
                    cols = max(len(tab.headers), max((len(r) for r in tab.rows), default=1), 1)
                    table = doc.add_table(rows=1 + len(tab.rows), cols=cols)
                    if tab.headers:
                        for i, h in enumerate(tab.headers[:cols]):
                            table.rows[0].cells[i].text = str(h)
                    for ri, row in enumerate(tab.rows):
                        for ci, cell in enumerate(row[:cols]):
                            table.rows[ri + 1].cells[ci].text = str(cell)
                    doc.add_paragraph(f"Table {tab.number}: {tab.caption or tab.title}")
                else:
                    doc.add_paragraph(block.text)
            else:
                doc.add_paragraph(block.text)

    bm = manuscript.back_matter
    for heading, value in (
        ("Acknowledgments", bm.acknowledgments),
        ("Funding", bm.funding),
        ("Conflict of Interest", bm.conflict_of_interest),
        ("Ethics", bm.ethics),
        ("Data Availability", bm.data_availability),
        ("Supplementary Materials", bm.supplementary_materials),
    ):
        if value.strip():
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)

    if manuscript.references:
        doc.add_heading("References", level=1)
        for ref in manuscript.references:
            authors = ", ".join(ref.authors)
            year = f" ({ref.year})" if ref.year else ""
            doc.add_paragraph(f"[{ref.order}] {authors}. {ref.title or ''}{year}.")

    note = doc.add_paragraph(tpl.warning)
    if note.runs:
        note.runs[0].italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
